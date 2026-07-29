import re
import json
import logging
import os
import time
from datetime import datetime
from pathlib import Path

import openpyxl
import pdfplumber
import pytesseract
import subprocess
from pdf2image import convert_from_path, pdfinfo_from_path
from PIL import Image
from docx import Document as DocxDocument

import ocr_engine

logging.basicConfig(level=logging.INFO, format="%(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Chiều dài khổ A4 theo inch (để tính DPI render vừa đủ, tránh ảnh quá lớn)
_A4_LONG_EDGE_IN = 11.69


def _effective_pdf_dpi(base_dpi, max_long_edge):
    """DPI render PDF sao cho cạnh dài ảnh ~ max_long_edge (không vượt base_dpi)."""
    dpi_fit = int(max_long_edge / _A4_LONG_EDGE_IN)
    return min(base_dpi, max(120, dpi_fit))


SUPPORTED_EXT = {
    ".pdf",
    ".docx",
    ".doc",
    ".xlsx",
    ".xls",
    ".pptx",
    ".ppt",
    ".md",
    ".txt",
}


# ================== UTIL ==================

def auto_rotate(image_path):
    try:
        img = Image.open(image_path)
        osd = pytesseract.image_to_osd(img, output_type=pytesseract.Output.DICT)
        angle = osd.get("rotate", 0)
        if angle:
            img.rotate(-angle, expand=True).save(image_path)
        return image_path
    except Exception:
        return image_path


def is_pdf_scan(pdf_path):
    try:
        total_chars = 0
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    total_chars += len(text.strip())
                    if total_chars >= 100:
                        return False
        return True
    except Exception as exc:
        logger.warning("Unable to read PDF text for scan check: %s", exc)
        return True


# ================== UTIL ==================

def _select_scan_pages(pdf_path, full_page_threshold, head_pages, tail_pages, dpi=200, max_long_edge=None):
    """Trả về danh sách (page_number, PIL.Image) các trang cần OCR.

    - File <= full_page_threshold trang (hoặc head/tail=None): render & OCR toàn bộ.
    - File lớn hơn: chỉ render head_pages trang đầu + tail_pages trang cuối (nhanh hơn).
    - dpi: độ phân giải render PDF; nếu có max_long_edge thì tự hạ dpi cho vừa (~ảnh chụp).
    """
    if max_long_edge:
        dpi = _effective_pdf_dpi(dpi, max_long_edge)
        logger.info("Render PDF dpi=%d (target cạnh dài ~%d px)", dpi, max_long_edge)
    if head_pages is None or tail_pages is None:
        pages = convert_from_path(pdf_path, dpi=dpi)
        return [(i + 1, page) for i, page in enumerate(pages)]

    total = pdfinfo_from_path(pdf_path)["Pages"]
    if total <= full_page_threshold:
        pages = convert_from_path(pdf_path, dpi=dpi)
        return [(i + 1, page) for i, page in enumerate(pages)]

    # Trang đầu: 1..head_pages ; Trang cuối: (total-tail_pages+1)..total (tránh chồng lấn)
    tail_start = max(head_pages + 1, total - tail_pages + 1)
    head_imgs = convert_from_path(pdf_path, dpi=dpi, first_page=1, last_page=head_pages)
    tail_imgs = convert_from_path(pdf_path, dpi=dpi, first_page=tail_start, last_page=total)
    selected = [(i, img) for i, img in enumerate(head_imgs, start=1)]
    selected += [(i, img) for i, img in enumerate(tail_imgs, start=tail_start)]
    logger.info("PDF %s trang -> chỉ OCR trang đầu 1-%d và cuối %d-%d", total, head_pages, tail_start, total)
    return selected


def convert_pdf_scan(pdf_path, engine, output_dir, *, full_page_threshold=5, head_pages=None, tail_pages=None, dpi=200, max_long_edge=None):
    """OCR các trang scan của PDF bằng `engine` (có hàm ocr_image)."""
    output = []
    os.makedirs(output_dir, exist_ok=True)
    for page_no, page in _select_scan_pages(
        pdf_path, full_page_threshold, head_pages, tail_pages, dpi=dpi, max_long_edge=max_long_edge,
    ):
        img_path = Path(output_dir) / f"temp_{page_no}.jpg"
        page.save(img_path, "JPEG", quality=92)
        # PDF render thường đúng chiều; bỏ OSD tesseract (~1–2s/trang), không ảnh hưởng chất lượng OCR

        t0 = time.time()
        text = engine.ocr_image(str(img_path))
        logger.info("OCR trang %s (%dx%d): %.1fs", page_no, page.size[0], page.size[1], time.time() - t0)

        # Hậu xử lý: Loại bỏ các bảng trống hoặc chỉ chứa tiêu đề bảng sửa đổi
        text = clean_empty_tables(text)

        # Định dạng markdown cho subtitle và text dựa trên quy luật phổ biến của OCR
        text = format_markdown_elements(text)

        output.append(f"\n## Page {page_no}\n{text}")
    return "\n".join(output)

def format_markdown_elements(text):
    # \"\"\"Hỗ trợ định dạng subtitle và text theo chuẩn markdown.\"\"\"
    lines = text.split('\n')
    formatted_lines = []
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            formatted_lines.append(line)
            continue
            
        # Nhận diện Subtitle: Thường là chữ in hoa, hoặc có số thứ tự La Mã/Chữ cái lớn ở đầu
        # Ví dụ: IV/ ĐỊNH NGHĨA..., A. QUY ĐỊNH CHUNG...
        is_subtitle = False
        if re.match(r'^[IVXLC]+\/.*', stripped) or \
           re.match(r'^[A-Z]\..*', stripped) or \
           (stripped.isupper() and len(stripped) < 100):
            is_subtitle = True
            
        if is_subtitle:
            # Nếu chưa có dấu # thì thêm vào để làm subtitle (H3)
            if not stripped.startswith('#'):
                formatted_lines.append(f"### {stripped}")
            else:
                formatted_lines.append(line)
        else:
            # Giữ nguyên text bình thường
            formatted_lines.append(line)
            
    return "\n".join(formatted_lines)

def clean_empty_tables(text):
    """Loại bỏ các bảng HTML không chứa dữ liệu thực tế hoặc chỉ chứa tiêu đề bảng sửa đổi."""
    # Tìm tất cả các khối <table>...</table>
    tables = re.findall(r'<table.*?>.*?</table>', text, flags=re.DOTALL | re.IGNORECASE)
    for table in tables:
        # Lấy nội dung TEXT của tất cả các ô <td>
        cells = re.findall(r'<td.*?>\s*(.*?)\s*</td>', table, flags=re.DOTALL | re.IGNORECASE)
        
        # Làm sạch nội dung từng ô (loại bỏ tag HTML bên trong ô và khoảng trắng)
        cleaned_cells = [re.sub(r'<.*?>', '', cell).strip() for cell in cells]
        
        # Loại bỏ các từ khóa tiêu đề bảng sửa đổi khỏi danh sách ô để kiểm tra dữ liệu thực
        keywords = {'lần', 'số yêu cầu', 'sửa đổi', 'nội dung', 'ngày', 'người', 'revision', 'date', 'description', 'author'}
        
        # Một ô được coi là có dữ liệu nếu nó không trống và không phải là từ khóa tiêu đề
        has_real_data = False
        for cell_text in cleaned_cells:
            if cell_text:
                # Kiểm tra xem cell_text có phải chỉ chứa từ khóa tiêu đề không
                is_header = any(kw in cell_text.lower() for kw in keywords)
                if not is_header:
                    has_real_data = True
                    break
        
        # Nếu bảng không có dữ liệu thực tế (chỉ có ô trống hoặc tiêu đề) thì xóa
        if not has_real_data:
            text = text.replace(table, "")
            
    # Xử lý trường hợp Page trống sau khi xóa bảng (nếu cần)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def convert_pdf_text(pdf_path):
    lines = []
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text:
                lines.append(f"\n## Page {i}\n{text}")
    return "\n".join(lines)


def convert_docling(file_path, converter=None):
    """Chuyển file sang markdown bằng docling. Lazy-load converter nếu chưa truyền."""
    if converter is None:
        from docling.document_converter import DocumentConverter
        converter = DocumentConverter()
    return converter.convert(file_path).document.export_to_markdown()


def convert_docx(file_path):
    doc = DocxDocument(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def convert_doc_to_docx(doc_path):
    dest = Path(doc_path).with_suffix(".docx")
    if dest.exists():
        return dest
    try:
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(dest.parent),
                str(doc_path),
            ],
            check=True,
            capture_output=True,
        )
    except subprocess.CalledProcessError as err:
        raise RuntimeError(f"Unable to convert {doc_path} to docx: {err.stderr.decode(errors='ignore')}" if err.stderr else f"Conversion failed for {doc_path}")
    if not dest.exists():
        raise RuntimeError(f"Converted file missing: {dest}")
    return dest


def convert_xlsx(file_path):
    wb = openpyxl.load_workbook(file_path)
    content = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        content.append(f"\n## Sheet: {sheet}")
        for row in ws.iter_rows(values_only=True):
            content.append(" | ".join(str(c) if c else "" for c in row))
    return "\n".join(content)


def save_markdown(content, file_path):
    output_path = Path(file_path).with_suffix(".md")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)


def convert_file(file_path, output_dir, engine, results):
    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".doc":
            file_path = convert_doc_to_docx(file_path)
            ext = ".docx"
        if ext == ".pdf":
            if is_pdf_scan(file_path):
                content = convert_pdf_scan(file_path, engine, output_dir)
            else:
                try:
                    content = convert_docling(file_path)
                except Exception as exc:
                    logger.warning("Docling lỗi cho %s, fallback pdfplumber: %s", file_path, exc)
                    content = convert_pdf_text(file_path)
        elif ext == ".docx":
            content = convert_docx(file_path)
        elif ext in {".xlsx", ".xls"}:
            content = convert_xlsx(file_path)
        else:
            content = convert_docling(file_path)
        if not content:
            raise ValueError("Empty content")
        save_markdown(content, file_path)
        results["success"].append(str(file_path))
    except Exception as exc:
        import traceback
        traceback.print_exc()
        logger.warning(f"Failed: {file_path} - {exc}")
        results["failed"].append({"file": str(file_path), "error": str(exc)})


def process_folder(input_dir, output_dir, engine):
    os.makedirs(output_dir, exist_ok=True)
    files = [
        os.path.join(root, f)
        for root, _, fs in os.walk(input_dir)
        for f in fs
        if Path(f).suffix.lower() in SUPPORTED_EXT
    ]
    results = {
        "timestamp": datetime.now().isoformat(),
        "total": len(files),
        "success": [],
        "failed": [],
    }
    for file_path in files:
        # Kiểm tra xem thư mục chứa file đã có file .md chưa, nếu có thì bỏ qua
        parent_dir = os.path.dirname(file_path)
        if any(f.lower().endswith('.md') for f in os.listdir(parent_dir)):
            logger.info(f"Skipping (Markdown exists): {os.path.basename(file_path)}")
            continue

        # if "Bộ quy tắc ứng xử" in file_path:
        logger.info(f"Processing: {os.path.basename(file_path)}")
        convert_file(file_path, output_dir, engine, results)
    summary_path = Path(output_dir) / "summary.json"
    with open(summary_path, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    logger.info(f"Done. Success: {len(results['success'])}")


def main():
    os.environ["CUDA_VISIBLE_DEVICES"] = "0"
    engine = ocr_engine.get_engine(device="0")
    for folder in [
        # "regulation", 
        # "procedures", 
        # "forms", 
        # "departments", 
        "guides"]:
        logger.info(f"--- Processing folder: {folder} ---")
        input_dir = f"/home/chatbot/AIPT_OCR/input"
        output_dir = input_dir
        process_folder(input_dir, output_dir, engine)

    if hasattr(engine, "close"):
        engine.close()


if __name__ == "__main__":
    main()


